# innovent_parser1.py  (DROP-IN MODULE)
from __future__ import annotations

import io
import os
import re
import shutil
import tempfile
import logging
from typing import Any

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from docx import Document  # python-docx

logger = logging.getLogger(__name__)


# -----------------------------
# Helpers
# -----------------------------

def _norm(s: str) -> str:
    return re.sub(r"[ \t]+", " ", (s or "")).strip()

def _safe_sheet_name(name: str) -> str:
    bad = r'[:\\/?*\[\]]'
    name = re.sub(bad, "-", (name or "")).strip()
    return (name[:31] or "Sheet1")


# -----------------------------
# DOC/DOCX -> DOCX bytes
# -----------------------------

def _try_read_docx_tables_and_text(docx_bytes: bytes) -> tuple[str, list[list[str]]]:
    """
    Attempt to load bytes as .docx using python-docx.
    Returns (full_text, flattened_table_rows).
    Raises if not a valid docx.
    """
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp_path = tmp.name

    try:
        doc = Document(tmp_path)
        full_text = "\n".join(_norm(p.text) for p in doc.paragraphs if _norm(p.text))

        rows: list[list[str]] = []
        for t in doc.tables:
            for r in t.rows:
                cells = [_norm(c.text) for c in r.cells]
                if any(cells):
                    rows.append(cells)

        return full_text, rows
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass


def _convert_doc_bytes_to_docx_bytes(doc_bytes: bytes) -> bytes:
    """
    Convert legacy .doc bytes to .docx bytes using LibreOffice.
    Requires 'soffice' to be installed.
    """
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("LibreOffice (soffice) not installed; cannot convert .doc. Upload .docx or install soffice.")

    with tempfile.TemporaryDirectory() as td:
        in_path = os.path.join(td, "input.doc")
        with open(in_path, "wb") as f:
            f.write(doc_bytes)

        import subprocess
        proc = subprocess.run(
            [soffice, "--headless", "--convert-to", "docx", "--outdir", td, in_path],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True
        )
        if proc.returncode != 0:
            raise RuntimeError(f"soffice .doc->.docx failed: {proc.stderr.strip() or proc.stdout.strip()}")

        # LibreOffice usually outputs input.docx
        out_path = os.path.join(td, "input.docx")
        if not os.path.exists(out_path):
            # fallback: first produced docx
            candidates = [p for p in os.listdir(td) if p.lower().endswith(".docx")]
            if not candidates:
                raise RuntimeError("soffice conversion produced no .docx file.")
            out_path = os.path.join(td, candidates[0])

        with open(out_path, "rb") as f:
            return f.read()


def _read_word_bytes(file_bytes: bytes) -> tuple[str, list[list[str]]]:
    """
    Works with either .docx or .doc bytes (auto-detect by trying docx first).
    """
    try:
        return _try_read_docx_tables_and_text(file_bytes)
    except Exception:
        # assume it's legacy .doc and convert
        docx_bytes = _convert_doc_bytes_to_docx_bytes(file_bytes)
        return _try_read_docx_tables_and_text(docx_bytes)


# -----------------------------
# Innovent extraction rules (your mapping)
# -----------------------------

def _extract_tag_model_from_table_rows(rows: list[list[str]]) -> tuple[str, str]:
    """
    Find header row containing Tag & Model, then take next row values:
      Tag | Qty | Model | Dims
    """
    for i in range(len(rows) - 1):
        header = [c.lower() for c in rows[i]]
        if any(c == "tag" for c in header) and any(c == "model" for c in header):
            for j in range(i + 1, min(i + 5, len(rows))):
                r = rows[j]
                if len(r) >= 3 and _norm(r[0]).lower() != "tag" and _norm(r[0]) and _norm(r[2]):
                    return _norm(r[0]), _norm(r[2])
    return "", ""


def _extract_section(full_text: str, title: str, stop_titles: list[str]) -> str:
    t = full_text or ""
    stops = "|".join(re.escape(s) for s in stop_titles) if stop_titles else r"\Z"

    m = re.search(
        rf"(?:^|\n){re.escape(title)}\s*\n(.*?)(?=\n(?:{stops})\s*\n|\Z)",
        t,
        flags=re.IGNORECASE | re.DOTALL
    )
    return m.group(1).strip() if m else ""


def parse_innovent_word_bytes(file_bytes: bytes) -> dict[str, Any]:
    full_text, table_rows = _read_word_bytes(file_bytes)

    tag, model = _extract_tag_model_from_table_rows(table_rows)

    # Sections per your requirement
    construction = _extract_section(full_text, "Construction", ["Components", "Comments"])
    components  = _extract_section(full_text, "Components",  ["Comments"])
    comments    = _extract_section(full_text, "Comments",    ["Innovent", "Innovent, LLC"])

    main_desc_parts: list[str] = []
    if construction:
        main_desc_parts.append("Construction")
        main_desc_parts.append(construction)
    if components:
        main_desc_parts.append("Components")
        main_desc_parts.append(components)

    main_desc = "\n\n".join(main_desc_parts).strip()

    return {
        "manufacturer": "Innovent",
        "header": {},
        "lines": [{
            "model_code": model,
            "tagging": tag,
            "tags": [tag] if tag else [""],
            "qty": 1,
            "list_each": 0.0,
            "total_list": 0.0,
            "options": [
                {"feature": "Equipment", "desc": main_desc or f"{tag} | {model}".strip(" |"), "add_price": 0.0},
                {"feature": "Include", "desc": comments.strip(), "add_price": 0.0} if comments.strip() else {},
            ]
        }]
    }


# -----------------------------
# Workbook writer (same vibe)
# -----------------------------

def _write_one_sheet(ws, *, manufacturer: str, equipment: str, model_code: str, tags: list[str], options: list[dict[str, Any]], output_type: str) -> None:
    output_type = (output_type or "all_in_one").strip().lower()

    # Row 1 headers A-F
    top_headers = ["Equipment","Manufacturer","Model","Part Number","Description (Not Overwritten)","Notes (Not Overwritten)"]
    for col, h in enumerate(top_headers, start=1):
        ws.cell(row=1, column=col, value=h)

    model_short = model_code.split("-")[0] if model_code else ""
    ws.cell(row=2, column=1, value=equipment)
    ws.cell(row=2, column=2, value=manufacturer)
    ws.cell(row=2, column=3, value=model_short)
    ws.cell(row=2, column=4, value=model_code)

    tags_label = ", ".join(t for t in (tags or []) if t) if tags else ""

    # Main desc in E2 from Equipment option only
    main_desc = ""
    include_descs: list[str] = []
    for opt in (options or []):
        if not isinstance(opt, dict):
            continue
        feat = (opt.get("feature") or "").strip().lower()
        if feat == "equipment":
            main_desc = (opt.get("desc") or "").strip()
        elif feat == "include":
            d = (opt.get("desc") or "").strip()
            if d:
                include_descs.append(d)

    ws.cell(row=2, column=5, value="\n\n".join([x for x in [tags_label, main_desc] if x]).strip())
    ws.cell(row=2, column=6, value="")

    # Row 3 option headers starting at C
    opt_headers = [
        "Tag","Part Number","Feature","Description","Qty","List Price","LP Ext.","Buy Mult.","Net Price","Markup","Margin","Sell Price",
        "Weight","Freight","Fr. Multi.","Alignment","Subtotal","Option Price"
    ]
    start_col = 3
    for i, h in enumerate(opt_headers):
        ws.cell(row=3, column=start_col + i, value=h)

    def z(): return 0
    row_idx = 4
    qty = 1

    # Your requirement: comments go in lower table as Include
    # We'll do it for both output types (safe); Innovent is “all_in_one” anyway.
    for tag in (tags or [""]):
        for d in include_descs:
            ws.cell(row=row_idx, column=3, value=tag)
            ws.cell(row=row_idx, column=4, value="")
            ws.cell(row=row_idx, column=5, value="Include")
            ws.cell(row=row_idx, column=6, value=d)

            ws.cell(row=row_idx, column=7, value=qty)
            ws.cell(row=row_idx, column=8, value=0.0)
            ws.cell(row=row_idx, column=9, value=0.0)

            ws.cell(row=row_idx, column=10, value=z())
            ws.cell(row=row_idx, column=11, value=z())
            ws.cell(row=row_idx, column=12, value=z())
            ws.cell(row=row_idx, column=13, value=z())
            ws.cell(row=row_idx, column=14, value=z())
            ws.cell(row=row_idx, column=15, value=z())
            ws.cell(row=row_idx, column=16, value=z())
            ws.cell(row=row_idx, column=17, value=z())
            ws.cell(row=row_idx, column=18, value=z())
            ws.cell(row=row_idx, column=19, value=z())
            ws.cell(row=row_idx, column=20, value=0.0)

            row_idx += 1

    # formatting
    header_fill = PatternFill("solid", fgColor="D9EAD3")
    top_header_font = Font(bold=True, italic=True)
    opt_header_font = Font(bold=True)

    for c in range(1, 7):
        cell = ws.cell(row=1, column=c)
        cell.font = top_header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for c in range(3, 21):
        cell = ws.cell(row=3, column=c)
        cell.font = opt_header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for r in range(2, max(row_idx, 4)):
        for c in range(1, 21):
            ws.cell(row=r, column=c).alignment = Alignment(vertical="top", wrap_text=True)

    widths = {1:18,2:18,3:16,4:28,5:40,6:55,7:8,8:12,9:12,10:10,11:12,12:10,13:10,14:12,15:10,16:10,17:10,18:10,19:12,20:12}
    for col, w in widths.items():
        ws.column_dimensions[get_column_letter(col)].width = w

    ws.freeze_panes = "A4"


def write_innovent_workbook(parsed: dict[str, Any], output_type: str = "all_in_one") -> Workbook:
    wb = Workbook()
    wb.remove(wb.active)

    manufacturer = parsed.get("manufacturer", "Innovent")
    lines = parsed.get("lines", []) or []
    if not lines:
        ws = wb.create_sheet("Innovent Import")
        _write_one_sheet(ws, manufacturer=manufacturer, equipment="Air Handler", model_code="", tags=[""], options=[], output_type=output_type)
        return wb

    for ln in lines:
        model_code = ln.get("model_code", "") or ""
        tags = ln.get("tags", [""]) or [""]
        options = [o for o in (ln.get("options", []) or []) if isinstance(o, dict) and o]

        ws = wb.create_sheet(_safe_sheet_name(model_code or "Innovent Import"))
        _write_one_sheet(ws, manufacturer=manufacturer, equipment="Air Handler", model_code=model_code, tags=tags, options=options, output_type=output_type)

    return wb


# -----------------------------
# ✅ ROUTE-COMPATIBLE RUNNER
# -----------------------------

def convert_innovent_pdf_to_xlsx_bytes(
    pdf_bytes: bytes,
    job_name: str | None = None,
    output_type: str = "all_in_one",
) -> tuple[bytes, str]:
    """
    Keep same signature as your existing Flask premise:
      converter_fn(pdf_bytes=..., job_name=..., output_type=...)
    """
    job = (job_name or "").strip() or "job"
    parsed = parse_innovent_word_bytes(pdf_bytes)

    wb = write_innovent_workbook(parsed, output_type=output_type)

    out = io.BytesIO()
    wb.save(out)
    xlsx_bytes = out.getvalue()
    out.close()

    filename = f"{job}_innovent_{output_type}_template_output.xlsx".replace(" ", "_")
    return xlsx_bytes, filename

