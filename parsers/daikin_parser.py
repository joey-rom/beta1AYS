from __future__ import annotations

import io
import re
from typing import Any

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

import logging
logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────

def _safe_sheet_name(name: str) -> str:
    name = re.sub(r'[:\\/?*\[\]]', "-", (name or "")).strip()
    return (name[:31] or "Sheet1")


def _money(s: str) -> float:
    s = re.sub(r'[,$]', '', s or '').strip()
    try:
        return float(s)
    except ValueError:
        return 0.0


# ─────────────────────────────────────────────
# Description template
# Keys map to the Description column in Table 2,
# parsed as "KEY; value" — we extract the value part.
#
# Each entry is one of:
#   ('literal', 'text')          → output text as-is
#   ('blank',)                   → blank line
#   ('field', 'KEY', 'LABEL')    → "LABEL; <value from doc>"
#   ('field', 'KEY')             → "KEY; <value from doc>"  (label = key)
# ─────────────────────────────────────────────

DESCRIPTION_TEMPLATE = [
    ('field',   'UNIT TYPE'),
    ('blank',),
    ('literal', '=== COOLING CAPACITY ==='),
    ('field',   'UNIT TONS'),
    ('field',   'UNIT EER'),
    ('field',   'UNIT IPLV'),
    ('field',   'REFRIGERANT TYPE'),
    ('field',   'UNIT MAKEUP'),
    ('field',   'BASE'),
    ('field',   'CONSTRUCTION'),
    ('blank',),
    ('literal', '=== GENERAL ==='),
    ('field',   'COMPRESSOR SIZE'),
    ('field',   'POWER CONNECTION'),
    ('field',   'SWITCH OPTIONS'),
    ('field',   'AGENCY APPROVAL'),
    ('field',   'AHRI APPROVAL'),
    ('field',   'ASHRAE APPROVAL'),
    ('blank',),
    ('literal', '=== EVAPORATOR ==='),
    ('field',   'EVAPORATOR TYPE/SIZE'),
    ('field',   'TUBE MATERIAL'),
    ('field',   'HEAD CONFIGURATION'),
    ('field',   'WATERSIDE PRESSURE'),
    ('field',   'EVAPORATOR INSULATION'),
    ('field',   'PERCENT OF GLYCOL'),
    ('blank',),
    ('literal', '=== CONDENSER ==='),
    ('field',   'FAN DESIGN TYPE'),
    ('field',   'GUARDS'),
    ('field',   'CONDENSER COIL FINS'),
    ('blank',),
    ('literal', '=== ELECTRICAL ==='),
    ('field',   'FAN NUMBER'),
    ('field',   'VOLTAGE'),
    ('field',   'STARTER TYPE/FILTER'),
    ('field',   'PHASE VOLTAGE'),
    ('field',   'MCA 1',  'MCA'),
    ('field',   'MOCP 1', 'MOCP'),
    ('field',   'GROUND FAULT'),
    ('blank',),
    ('field',   'COMMUNICATION'),
    ('field',   'DISPLAY OPTION'),
    ('blank',),
    ('literal', '=== EQUIPMENT START-UP AND FIELD SERVICES ==='),
    ('literal', '* Equipment start-up (1 day)'),
    ('literal', '* Owner training (1 day)'),
    ('blank',),
    ('literal', '=== EQUIPMENT WARRANTY ==='),
    ('literal', "* Manufacturer's first (1st) year __parts__ warranty from date of equipment "
                "start-up not to exceed eighteen (18) months from date of shipment, whichever "
                "occurs first (refrigerant and labor warranty not included)"),
    ('field',   '1ST YEAR WARRANTY'),
    ('field',   'EXT. UNIT WARRANTY'),
    ('field',   'REFRIGERANT WARRANTY'),
    ('blank',),
    ('blank',),
    ('literal', '=== NOT INCLUDED ==='),
    ('literal', 'Installation, rigging, storage, unloading, equipment commissioning, system '
                'commissioning, diagnostic service calls, refrigerant warranty, labor warranty, '
                'maintenance, coil cleaning, quarterly inspections, annual inspections, power or '
                'control wiring, piping, 20 mesh cleanable chilled water strainer upstream of '
                'evaporator, valves, rigging, or unloading.'),
]


def _build_description(option_rows: dict[str, str]) -> str:
    """Build the formatted description string from parsed option rows."""
    lines = []
    for entry in DESCRIPTION_TEMPLATE:
        if entry[0] == 'blank':
            lines.append('')
        elif entry[0] == 'literal':
            lines.append(entry[1])
        elif entry[0] == 'field':
            key   = entry[1]
            label = entry[2] if len(entry) > 2 else key
            # Strip trailing spaces from key for lookup
            val = option_rows.get(key, option_rows.get(key.strip(), ''))
            lines.append(f'{label}; {val}' if val else f'{label};')
    return '\n'.join(lines)


# ─────────────────────────────────────────────
# DOCX PARSER
# ─────────────────────────────────────────────

def parse_docx_bytes(file_bytes: bytes, filename: str = "") -> dict[str, Any]:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))

    # ── Header from Table 0 ─────────────────────
    # Row 1: Job Name | value | value | Date | date_value
    # Row 2: Model    | value | ...   | Software Version | ver
    # Row 3: Unit Tag | value | ...
    job_name  = ""
    model     = ""
    unit_tag  = ""
    quote_date = ""

    try:
        t0 = doc.tables[0]
        job_name   = t0.rows[1].cells[1].text.strip()
        quote_date = t0.rows[1].cells[4].text.strip()
        model      = t0.rows[2].cells[1].text.strip()
        unit_tag   = t0.rows[3].cells[1].text.strip()
    except (IndexError, AttributeError):
        pass

    header = {
        "quote_number": model,
        "quote_date":   quote_date,
        "job_name":     job_name,
        "unit_tag":     unit_tag,
        "customer":     "",
        "requested_ship_date": "",
    }

    # ── Option rows from Table 2 ─────────────────
    # Each data row: [Code Item, Code Value, Description, List Price]
    # Description format: "KEY; Human readable value"
    # We also need list price and qty from the summary rows at the bottom.

    option_rows: dict[str, str] = {}  # KEY -> human-readable value
    option_list: list[dict[str, Any]] = []  # for shopping list
    list_each   = 0.0
    qty         = 1
    total_list  = 0.0

    try:
        t2 = doc.tables[2]
        for row in t2.rows:
            cells = [c.text.strip() for c in row.cells]
            code_item  = cells[0] if len(cells) > 0 else ''
            code_value = cells[1] if len(cells) > 1 else ''
            desc_cell  = cells[2] if len(cells) > 2 else ''
            price_cell = cells[3] if len(cells) > 3 else ''

            # Summary rows at the bottom
            if 'List Each:' in desc_cell:
                list_each = _money(price_cell)
                continue
            if 'Quantity:' in desc_cell:
                m = re.search(r'x\s*(\d+)', price_cell)
                if m: qty = int(m.group(1))
                continue
            if 'Total Ext List:' in desc_cell:
                total_list = _money(price_cell)
                continue

            # Option rows: "KEY; Value"
            if ';' in desc_cell:
                key, _, val = desc_cell.partition(';')
                key = key.strip()
                val = val.strip()
                option_rows[key] = val
                price = _money(price_cell)
                option_list.append({
                    "code_item":   code_item,
                    "code_value":  code_value,
                    "key":         key,
                    "desc":        val,
                    "full_desc":   desc_cell,
                    "unit_price":  price,
                    "total_price": price,
                    "qty":         1,
                })

    except (IndexError, AttributeError):
        pass

    if not option_rows:
        return {"manufacturer": "Daikin", "header": header, "lines": []}

    description = _build_description(option_rows)
    equipment   = option_rows.get('UNIT TYPE', 'Air Cooled Scroll Chiller')

    line = {
        "model_code":  model,
        "description": description,
        "equipment":   equipment,
        "tagging":     unit_tag,
        "tags":        [unit_tag] if unit_tag else [],
        "qty":         qty,
        "list_each":   list_each,
        "total_list":  total_list,
        "options":     option_list,
    }

    return {"manufacturer": "Daikin", "header": header, "lines": [line]}


# ─────────────────────────────────────────────
# XLSX WRITER
# ─────────────────────────────────────────────

def _write_one_sheet(ws, *, manufacturer, equipment, model_code,
                     description, tags, header, qty, list_each) -> None:

    top_headers = ["Equipment", "Manufacturer", "Model", "Part Number",
                   "Description (Not Overwritten)", "Notes (Not Overwritten)"]
    for col, h in enumerate(top_headers, start=1):
        ws.cell(row=1, column=col, value=h)

    quote_no   = header.get("quote_number", "")
    job_name   = header.get("job_name", "")
    quote_date = header.get("quote_date", "")
    unit_tag   = header.get("unit_tag", "")

    model_short = model_code.split("-")[0] if model_code else model_code

    ws.cell(row=2, column=1, value=equipment)
    ws.cell(row=2, column=2, value=manufacturer)
    ws.cell(row=2, column=3, value=model_short)
    ws.cell(row=2, column=4, value=model_code)
    ws.cell(row=2, column=5, value=description)

    notes_parts = []
    if quote_no:   notes_parts.append(f"Model: {quote_no}")
    if job_name:   notes_parts.append(f"Job: {job_name}")
    if quote_date: notes_parts.append(f"Date: {quote_date}")
    ws.cell(row=2, column=6, value="\n".join(notes_parts))

    opt_headers = [
        "Tag", "Part Number", "Feature", "Description",
        "Qty", "List Price", "LP Ext.", "Buy Mult.", "Net Price",
        "Markup", "Margin", "Sell Price", "Weight", "Freight",
        "Fr. Multi.", "Alignment", "Subtotal", "Option Price",
    ]
    for i, h in enumerate(opt_headers):
        ws.cell(row=3, column=3 + i, value=h)

    # Formatting
    header_fill = PatternFill("solid", fgColor="D9EAD3")
    top_font    = Font(bold=True, italic=True, name="Arial")
    opt_font    = Font(bold=True, name="Arial")
    body_font   = Font(name="Arial")

    for c in range(1, 7):
        cell = ws.cell(row=1, column=c)
        cell.font = top_font; cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for c in range(3, 21):
        cell = ws.cell(row=3, column=c)
        cell.font = opt_font; cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for r in range(2, 4):
        for c in range(1, 21):
            cell = ws.cell(row=r, column=c)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.font = body_font

    widths = {1:28, 2:18, 3:16, 4:20, 5:55, 6:45,
              7:8,  8:12, 9:12, 10:10, 11:12, 12:10,
              13:10, 14:12, 15:10, 16:10, 17:10, 18:10, 19:12, 20:12}
    for col, w in widths.items():
        ws.column_dimensions[get_column_letter(col)].width = w

    ws.freeze_panes = "A4"


def _write_shopping_list_sheet(ws, *, manufacturer, equipment, model_code,
                               tags, header, qty, list_each, options) -> None:
    """Shopping list: simple header row + one include line per option."""
    top_headers = ["Equipment", "Manufacturer", "Model", "Part Number",
                   "Description (Not Overwritten)", "Notes (Not Overwritten)"]
    for col, h in enumerate(top_headers, start=1):
        ws.cell(row=1, column=col, value=h)

    quote_no   = header.get("quote_number", "")
    job_name   = header.get("job_name", "")
    quote_date = header.get("quote_date", "")
    unit_tag   = header.get("unit_tag", "")
    model_short = model_code.split("-")[0] if model_code else model_code

    # Row 2: just the basic equipment header — no big description block
    ws.cell(row=2, column=1, value=equipment)
    ws.cell(row=2, column=2, value=manufacturer)
    ws.cell(row=2, column=3, value=model_short)
    ws.cell(row=2, column=4, value=model_code)
    ws.cell(row=2, column=5, value=equipment)

    notes_parts = []
    if quote_no:   notes_parts.append(f"Model: {quote_no}")
    if job_name:   notes_parts.append(f"Job: {job_name}")
    if quote_date: notes_parts.append(f"Date: {quote_date}")
    ws.cell(row=2, column=6, value="\n".join(notes_parts))

    opt_headers = [
        "Tag", "Part Number", "Feature", "Description",
        "Qty", "List Price", "LP Ext.", "Buy Mult.", "Net Price",
        "Markup", "Margin", "Sell Price", "Weight", "Freight",
        "Fr. Multi.", "Alignment", "Subtotal", "Option Price",
    ]
    for i, h in enumerate(opt_headers):
        ws.cell(row=3, column=3 + i, value=h)

    tag = ", ".join(tags) if tags else ""
    row_idx = 4
    for opt in (options or []):
        unit_price = float(opt.get("unit_price") or 0)
        opt_qty    = int(opt.get("qty") or 1)
        lp_ext     = unit_price * opt_qty
        full_desc  = opt.get("full_desc") or opt.get("desc", "")

        ws.cell(row=row_idx, column=3,  value=opt.get("code_item", ""))
        ws.cell(row=row_idx, column=4,  value=opt.get("code_value", ""))
        ws.cell(row=row_idx, column=5,  value="Include")
        ws.cell(row=row_idx, column=6,  value=full_desc)
        ws.cell(row=row_idx, column=7,  value=opt_qty)
        ws.cell(row=row_idx, column=8,  value=unit_price)
        ws.cell(row=row_idx, column=9,  value=lp_ext)
        for c in range(10, 21):
            ws.cell(row=row_idx, column=c, value=0)
        ws.cell(row=row_idx, column=20, value=unit_price)
        row_idx += 1

    # Formatting
    header_fill = PatternFill("solid", fgColor="D9EAD3")
    top_font    = Font(bold=True, italic=True, name="Arial")
    opt_font    = Font(bold=True, name="Arial")
    body_font   = Font(name="Arial")

    for c in range(1, 7):
        cell = ws.cell(row=1, column=c)
        cell.font = top_font; cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for c in range(3, 21):
        cell = ws.cell(row=3, column=c)
        cell.font = opt_font; cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for r in range(2, max(row_idx, 4)):
        for c in range(1, 21):
            cell = ws.cell(row=r, column=c)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.font = body_font

    widths = {1:28, 2:18, 3:16, 4:20, 5:55, 6:45,
              7:8,  8:12, 9:12, 10:10, 11:12, 12:10,
              13:10, 14:12, 15:10, 16:10, 17:10, 18:10, 19:12, 20:12}
    for col, w in widths.items():
        ws.column_dimensions[get_column_letter(col)].width = w

    ws.freeze_panes = "A4"


def write_template_workbook(parsed: dict[str, Any], output_type: str = "all_in_one") -> Workbook:
    wb = Workbook()
    wb.remove(wb.active)

    manufacturer = parsed.get("manufacturer", "")
    header       = parsed.get("header", {})
    lines        = parsed.get("lines", [])

    if not lines:
        ws = wb.create_sheet("Import")
        _write_one_sheet(ws, manufacturer=manufacturer, equipment="",
                         model_code="", description="", tags=[],
                         header=header, qty=1, list_each=0.0)
        return wb

    sheet_names_seen: set[str] = set()
    for ln in lines:
        model_code  = ln.get("model_code", "") or ""
        description = ln.get("description", "") or ""
        tagging     = ln.get("tagging", "") or ""
        tags        = ln.get("tags") or ([tagging] if tagging else [])
        equipment   = ln.get("equipment", "")
        qty         = ln.get("qty", 1)
        list_each   = ln.get("list_each", 0.0)
        options     = ln.get("options", [])

        base_name = _safe_sheet_name(model_code or "Import")
        name = base_name
        n = 2
        while name in sheet_names_seen:
            suffix = f"_{n}"
            name = _safe_sheet_name(base_name[:31 - len(suffix)] + suffix)
            n += 1
        sheet_names_seen.add(name)

        ws = wb.create_sheet(title=name)

        if output_type == "shopping_list":
            _write_shopping_list_sheet(ws, manufacturer=manufacturer, equipment=equipment,
                                       model_code=model_code, tags=tags, header=header,
                                       qty=qty, list_each=list_each, options=options)
        else:
            _write_one_sheet(ws, manufacturer=manufacturer, equipment=equipment,
                             model_code=model_code, description=description,
                             tags=tags, header=header, qty=qty, list_each=list_each)

    return wb


# ─────────────────────────────────────────────
# PUBLIC API
# ─────────────────────────────────────────────

def convert_daikin_to_xlsx_bytes(
    pdf_bytes: bytes,
    filename: str = "",
    job_name: str | None = None,
    output_type: str | None = None,
    **kwargs,
) -> tuple[bytes, str]:
    parsed = parse_docx_bytes(pdf_bytes, filename=filename)

    if job_name:
        parsed["header"]["job_name"] = job_name.strip()

    wb = write_template_workbook(parsed, output_type=output_type or "all_in_one")

    out = io.BytesIO()
    wb.save(out)
    xlsx_bytes = out.getvalue()
    out.close()

    job = (parsed["header"].get("job_name") or
           re.sub(r'\.(docx|pdf)$', '', filename, flags=re.IGNORECASE) or
           "output")
    return xlsx_bytes, f"{job}_template_output.xlsx"


# ─────────────────────────────────────────────
# Quick test when run directly
# ─────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python3 daikin_parser.py <path_to_docx> [all_in_one|shopping_list]")
        sys.exit(1)
    path = sys.argv[1]
    output_type = sys.argv[2] if len(sys.argv) > 2 else "all_in_one"
    with open(path, "rb") as f:
        data = f.read()
    fname = path.split("/")[-1]
    xlsx_bytes, out_name = convert_daikin_to_xlsx_bytes(data, filename=fname, output_type=output_type)
    out_path = f"/mnt/user-data/outputs/{out_name}"
    with open(out_path, "wb") as f:
        f.write(xlsx_bytes)
    print(f"Written: {out_path}")