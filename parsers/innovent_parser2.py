# innovent_parser1.py  (DROP-IN MODULE — FIXED TAG/MODEL + PRESERVE NUMBERING)
from __future__ import annotations

import io
import os
import re
import shutil
import tempfile
import logging
from typing import Any

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from docx import Document  # python-docx

logger = logging.getLogger(__name__)


# -----------------------------
# Helpers
# -----------------------------

def _safe_sheet_name(name: str) -> str:
    bad = r'[:\\/?*\[\]]'
    name = re.sub(bad, "-", (name or "")).strip()
    return (name[:31] or "Sheet1")


def _norm(s: str) -> str:
    return re.sub(r"[ \t]+", " ", (s or "")).strip()


def _clean_quotes(s: str) -> str:
    # normalize fancy quotes to plain
    return (
        (s or "")
        .replace("“", '"').replace("”", '"')
        .replace("’", "'").replace("‘", "'")
        .strip()
    )


# -----------------------------
# Numbering reconstruction (Word lists)
# -----------------------------

def _paragraph_num_info(p) -> tuple[int | None, int | None]:
    """
    Returns (numId, ilvl) if paragraph is part of a numbered/bulleted list.
    Uses underlying XML (python-docx doesn't include list numbers in paragraph.text).
    """
    try:
        pPr = p._p.pPr
        if pPr is None or pPr.numPr is None:
            return None, None
        numId = pPr.numPr.numId.val if pPr.numPr.numId is not None else None
        ilvl = pPr.numPr.ilvl.val if pPr.numPr.ilvl is not None else None
        return int(numId) if numId is not None else None, int(ilvl) if ilvl is not None else None
    except Exception:
        return None, None


def _iter_paragraphs_with_numbering(doc: Document) -> list[str]:
    """
    Build paragraph text with reconstructed numbering prefixes like:
      1. line
      1.1 subline
    This is a pragmatic numbering renderer; it won’t perfectly match every Word style,
    but it reliably preserves “the numbers on the left” for typical numbered lists.
    """
    counters: dict[tuple[int, int], int] = {}  # (numId, ilvl) -> count
    last_key: tuple[int, int] | None = None

    out: list[str] = []
    for p in doc.paragraphs:
        txt = _clean_quotes(p.text or "")
        if not txt.strip():
            continue

        numId, ilvl = _paragraph_num_info(p)
        if numId is None or ilvl is None:
            # reset tracking when leaving lists
            last_key = None
            out.append(txt.strip())
            continue

        # increment this level, reset deeper levels for this numId
        key = (numId, ilvl)
        counters[key] = counters.get(key, 0) + 1

        # reset deeper levels
        for (n, lvl) in list(counters.keys()):
            if n == numId and lvl > ilvl:
                counters.pop((n, lvl), None)

        # build prefix: join levels 0..ilvl if present
        parts: list[str] = []
        for lvl in range(0, ilvl + 1):
            k2 = (numId, lvl)
            if k2 in counters:
                parts.append(str(counters[k2]))
            else:
                # if a parent level wasn't seen, don't invent it
                pass

        prefix = ".".join(parts) + ". " if parts else ""
        out.append(prefix + txt.strip())
        last_key = key

    return out


# -----------------------------
# DOC/DOCX -> DOCX bytes
# -----------------------------

def _try_read_docx_tables_and_text(docx_bytes: bytes) -> tuple[str, list[list[str]]]:
    """
    Load bytes as .docx using python-docx.
    Returns:
      - full_text (with numbering prefixes reconstructed)
      - flattened table rows
    """
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp_path = tmp.name

    try:
        doc = Document(tmp_path)

        # paragraphs with numbering preserved
        para_lines = _iter_paragraphs_with_numbering(doc)
        full_text = "\n".join(para_lines)

        # tables (when they are real Word tables)
        rows: list[list[str]] = []
        for t in doc.tables:
            for r in t.rows:
                cells = [_clean_quotes(c.text) for c in r.cells]
                cells = [_norm(c) for c in cells]
                if any(cells):
                    rows.append(cells)

        return full_text, rows
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass


def _convert_doc_bytes_to_docx_bytes(doc_bytes: bytes) -> bytes:
    """
    Convert legacy .doc bytes to .docx bytes using LibreOffice.
    Requires 'soffice' installed on the server.
    """
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError(
            "LibreOffice (soffice) not installed; cannot convert .doc. "
            "Upload .docx or install soffice."
        )

    with tempfile.TemporaryDirectory() as td:
        in_path = os.path.join(td, "input.doc")
        with open(in_path, "wb") as f:
            f.write(doc_bytes)

        import subprocess
        proc = subprocess.run(
            [soffice, "--headless", "--convert-to", "docx", "--outdir", td, in_path],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True
        )
        if proc.returncode != 0:
            raise RuntimeError(f"soffice .doc->.docx failed: {proc.stderr.strip() or proc.stdout.strip()}")

        out_path = os.path.join(td, "input.docx")
        if not os.path.exists(out_path):
            candidates = [p for p in os.listdir(td) if p.lower().endswith(".docx")]
            if not candidates:
                raise RuntimeError("soffice conversion produced no .docx file.")
            out_path = os.path.join(td, candidates[0])

        with open(out_path, "rb") as f:
            return f.read()


def _read_word_bytes(file_bytes: bytes) -> tuple[str, list[list[str]]]:
    """
    Works with either .docx or .doc bytes (auto-detect).
    """
    try:
        return _try_read_docx_tables_and_text(file_bytes)
    except Exception:
        docx_bytes = _convert_doc_bytes_to_docx_bytes(file_bytes)
        return _try_read_docx_tables_and_text(docx_bytes)


# -----------------------------
# Innovent extraction rules
# -----------------------------

def _extract_quote_number(full_text: str) -> str:
    m = re.search(r"Quotation\s*No\.\s*([^\n]+)", full_text or "", flags=re.IGNORECASE)
    return _norm(m.group(1)) if m else ""


def _extract_tag_qty_model_dims_from_table_rows(rows: list[list[str]]) -> tuple[str, int, str, str]:
    """
    If the doc uses real Word tables, this works.
    Header row contains: Tag | Qty | Model | Dims
    """
    for i in range(len(rows) - 1):
        header = [_norm(c).lower() for c in rows[i]]
        if ("tag" in header) and ("qty" in header) and ("model" in header) and ("dims" in header):
            # data row usually immediately below
            for j in range(i + 1, min(i + 6, len(rows))):
                r = rows[j]
                if len(r) < 4:
                    continue
                if _norm(r[0]).lower() == "tag":
                    continue
                tag = _norm(r[0])
                qty_s = _norm(r[1])
                model = _norm(r[2])
                dims = _norm(r[3])
                if tag and model:
                    try:
                        qty = int(re.sub(r"[^0-9]", "", qty_s) or "0")
                    except Exception:
                        qty = 0
                    return tag, qty, model, dims
    return "", 0, "", ""


def _extract_tag_qty_model_dims_from_text(full_text: str) -> tuple[str, int, str, str]:
    """
    Strong text-based parser that works even when the 'table' isn't a real table.
    Finds the header line containing Tag/Qty/Model/Dims then parses the next lines.

    Handles cases where Dims wraps to next line.
    """
    t = full_text or ""
    lines = [ln.rstrip() for ln in t.splitlines()]

    # find header line index
    hdr_idx = None
    for i, ln in enumerate(lines):
        l = ln.lower()
        if ("tag" in l) and ("qty" in l) and ("model" in l) and ("dims" in l):
            hdr_idx = i
            break
    if hdr_idx is None:
        return "", 0, "", ""

    # collect following non-empty lines until we hit "Price:" etc.
    stop_re = re.compile(r"^\s*(Price:|Terms:|Validity:|Shipment:|Add/deducts:)\b", re.IGNORECASE)
    data_lines: list[str] = []
    for ln in lines[hdr_idx + 1:]:
        if stop_re.search(ln):
            break
        if ln.strip():
            data_lines.append(ln)

        # usually first 1–3 lines contain the row (dims might wrap)
        if len(data_lines) >= 3:
            break

    if not data_lines:
        return "", 0, "", ""

    # Attempt parse using "2+ spaces" as columns
    def split_cols(s: str) -> list[str]:
        s = _clean_quotes(s)
        parts = re.split(r"\s{2,}", s.strip())
        return [p.strip() for p in parts if p.strip()]

    cols = split_cols(data_lines[0])

    # If the row wrapped, append continuation text to last col (dims)
    if len(data_lines) > 1 and cols and len(cols) < 4:
        # try combine first two lines then split again
        cols2 = split_cols(data_lines[0] + "  " + data_lines[1])
        if len(cols2) > len(cols):
            cols = cols2
        else:
            # otherwise treat line2 as dims continuation
            cols = cols + [" ".join(split_cols(data_lines[1]))]

    # Expect at least Tag, Qty, Model; Dims may be everything after
    if len(cols) < 3:
        # fallback: simple token parse
        tok = data_lines[0].split()
        if len(tok) >= 3:
            tag = tok[0]
            try:
                qty = int(tok[1])
            except Exception:
                qty = 0
            model = tok[2]
            dims = " ".join(tok[3:]).strip()
            if len(data_lines) > 1:
                dims = (dims + " " + " ".join(data_lines[1:])).strip()
            return _norm(tag), qty, _norm(model), _norm(dims)
        return "", 0, "", ""

    tag = cols[0]
    try:
        qty = int(re.sub(r"[^0-9]", "", cols[1]) or "0")
    except Exception:
        qty = 0
    model = cols[2]

    dims = ""
    if len(cols) >= 4:
        dims = " ".join(cols[3:]).strip()
    else:
        # if not present, try remaining lines
        if len(data_lines) > 1:
            dims = " ".join([ln.strip() for ln in data_lines[1:]]).strip()

    return _norm(tag), int(qty), _norm(model), _norm(dims)


def _extract_tag_qty_model_dims(rows: list[list[str]], full_text: str) -> tuple[str, int, str, str]:
    tag, qty, model, dims = _extract_tag_qty_model_dims_from_table_rows(rows)
    if tag and model:
        return tag, qty, model, dims
    return _extract_tag_qty_model_dims_from_text(full_text)


def _extract_section(full_text: str, title: str, stop_titles: list[str]) -> str:
    """
    Extract section body under a header line equal to title, preserving numbering prefixes
    (already present in full_text due to our numbering renderer).
    """
    t = full_text or ""
    stops = "|".join(re.escape(s) for s in stop_titles) if stop_titles else r"\Z"

    m = re.search(
        rf"(?:^|\n){re.escape(title)}\s*\n(.*?)(?=\n(?:{stops})\s*\n|\Z)",
        t,
        flags=re.IGNORECASE | re.DOTALL
    )
    return m.group(1).strip() if m else ""


# -----------------------------
# Parser (Word bytes -> structured dict)
# -----------------------------

def parse_innovent_word_bytes(file_bytes: bytes) -> dict[str, Any]:
    full_text, table_rows = _read_word_bytes(file_bytes)

    quote_no = _extract_quote_number(full_text)
    tag, qty, model, dims = _extract_tag_qty_model_dims(table_rows, full_text)

    # Sections per your requirement
    construction = _extract_section(full_text, "Construction", ["Components", "Comments"])
    components  = _extract_section(full_text, "Components",  ["Comments"])
    comments    = _extract_section(full_text, "Comments",    ["Innovent", "Innovent, LLC"])

    main_desc_parts: list[str] = []
    if construction:
        main_desc_parts.append("Construction")
        main_desc_parts.append(construction)
    if components:
        main_desc_parts.append("Components")
        main_desc_parts.append(components)

    main_desc = "\n\n".join([p for p in main_desc_parts if _norm(p)]).strip()

    options: list[dict[str, Any]] = []
    options.append({
        "feature": "Equipment",
        "desc": main_desc or f"{tag} | {model}".strip(" |"),
        "add_price": 0.0
    })
    if comments.strip():
        options.append({
            "feature": "Include",
            "desc": comments.strip(),
            "add_price": 0.0
        })

    return {
        "manufacturer": "Innovent",
        "header": {
            "quote_number": quote_no,
            "dims": dims,
        },
        "lines": [{
            "model_code": model,
            "tagging": tag,
            "tags": [tag] if tag else [""],
            "qty": int(qty or 0),
            "list_each": 0.0,
            "total_list": 0.0,
            "options": options
        }]
    }


# -----------------------------
# Workbook writer (same vibe as your other converters)
# -----------------------------
def _write_one_sheet(
    ws,
    *,
    manufacturer: str,
    equipment: str,
    model_code: str,
    tags: list[str],
    header: dict[str, Any],
    line_qty: int,
    options: list[dict[str, Any]],
    output_type: str
) -> None:
    output_type = (output_type or "all_in_one").strip().lower()

    # Row 1 headers A-F
    top_headers = [
        "Equipment", "Manufacturer", "Model", "Part Number",
        "Description (Not Overwritten)", "Notes (Not Overwritten)"
    ]
    for col, h in enumerate(top_headers, start=1):
        ws.cell(row=1, column=col, value=h)

    # Row 2 values
    ws.cell(row=2, column=1, value=equipment)     # always "Air Handler"
    ws.cell(row=2, column=2, value=manufacturer)

    # ✅ Your requested mapping:
    # C2 (Model) = full model
    # D2 (Part Number) = blank
    ws.cell(row=2, column=3, value=model_code)
    ws.cell(row=2, column=4, value="")

    # Tag(s) + Qty for NOTES (not Description)
    tags_label = ", ".join(t for t in (tags or []) if t) if tags else ""
    qty_val = int(line_qty) if (line_qty and int(line_qty) > 0) else 0

    # Extract main desc + include descs from options
    main_desc = ""
    include_descs: list[str] = []
    for opt in (options or []):
        if not isinstance(opt, dict) or not opt:
            continue
        feat = (opt.get("feature") or "").strip().lower()
        if feat == "equipment":
            main_desc = (opt.get("desc") or "").strip()
        elif feat == "include":
            d = (opt.get("desc") or "").strip()
            if d:
                include_descs.append(d)

    # ✅ E2 (Description) = Construction/Components ONLY
    ws.cell(row=2, column=5, value=main_desc.strip())

    # ✅ F2 (Notes) = Tag + Qty + Quote # + Dims
    quote_no = _norm(str((header or {}).get("quote_number") or ""))
    dims = _norm(str((header or {}).get("dims") or ""))

    notes_lines: list[str] = []
    if tags_label:
        notes_lines.append(f"Tag: {tags_label}")
    if qty_val > 0:
        notes_lines.append(f"Qty: {qty_val}")
    if quote_no:
        notes_lines.append(f"Quote: {quote_no}")
    if dims:
        notes_lines.append(f"Dims: {dims}")

    ws.cell(row=2, column=6, value="\n".join(notes_lines).strip())

    # Row 3 option headers starting at C
    opt_headers = [
        "Tag", "Part Number", "Feature", "Description",
        "Qty", "List Price", "LP Ext.", "Buy Mult.", "Net Price", "Markup", "Margin", "Sell Price",
        "Weight", "Freight", "Fr. Multi.", "Alignment", "Subtotal", "Option Price"
    ]
    start_col = 3
    for i, h in enumerate(opt_headers):
        ws.cell(row=3, column=start_col + i, value=h)

    def z() -> int:
        return 0

    row_idx = 4
    option_qty = 1
    option_price = 0.0

    # Lower table: Include rows (Comments)
    for tag in (tags or [""]):
        for d in include_descs:
            ws.cell(row=row_idx, column=3, value=tag)
            ws.cell(row=row_idx, column=4, value="")           # Part Number
            ws.cell(row=row_idx, column=5, value="Include")    # Feature
            ws.cell(row=row_idx, column=6, value=d)            # Description

            ws.cell(row=row_idx, column=7, value=option_qty)
            ws.cell(row=row_idx, column=8, value=option_price)
            ws.cell(row=row_idx, column=9, value=option_price)

            ws.cell(row=row_idx, column=10, value=z())
            ws.cell(row=row_idx, column=11, value=z())
            ws.cell(row=row_idx, column=12, value=z())
            ws.cell(row=row_idx, column=13, value=z())
            ws.cell(row=row_idx, column=14, value=z())
            ws.cell(row=row_idx, column=15, value=z())
            ws.cell(row=row_idx, column=16, value=z())
            ws.cell(row=row_idx, column=17, value=z())
            ws.cell(row=row_idx, column=18, value=z())
            ws.cell(row=row_idx, column=19, value=z())
            ws.cell(row=row_idx, column=20, value=option_price)

            row_idx += 1

    # formatting
    header_fill = PatternFill("solid", fgColor="D9EAD3")
    top_header_font = Font(bold=True, italic=True)
    opt_header_font = Font(bold=True)

    for c in range(1, 7):
        cell = ws.cell(row=1, column=c)
        cell.font = top_header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for c in range(3, 21):
        cell = ws.cell(row=3, column=c)
        cell.font = opt_header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for r in range(2, max(row_idx, 4)):
        for c in range(1, 21):
            ws.cell(row=r, column=c).alignment = Alignment(vertical="top", wrap_text=True)

    widths = {
        1: 18, 2: 18, 3: 16, 4: 28, 5: 40, 6: 55,
        7: 8, 8: 12, 9: 12, 10: 10, 11: 12, 12: 10, 13: 10, 14: 12,
        15: 10, 16: 10, 17: 10, 18: 10, 19: 12, 20: 12
    }
    for col, w in widths.items():
        ws.column_dimensions[get_column_letter(col)].width = w

    ws.freeze_panes = "A4"


def write_innovent_workbook(parsed: dict[str, Any], output_type: str = "all_in_one") -> Workbook:
    wb = Workbook()
    wb.remove(wb.active)

    manufacturer = parsed.get("manufacturer", "Innovent")
    header = parsed.get("header", {}) or {}
    lines = parsed.get("lines", []) or []

    if not lines:
        ws = wb.create_sheet("Innovent Import")
        _write_one_sheet(
            ws,
            manufacturer=manufacturer,
            equipment="Air Handler",
            model_code="",
            tags=[""],
            header=header,
            line_qty=1,
            options=[],
            output_type=output_type
        )
        return wb

    sheet_names_seen: set[str] = set()
    for ln in lines:
        model_code = ln.get("model_code", "") or ""
        tags = ln.get("tags", [""]) or [""]
        options = [o for o in (ln.get("options", []) or []) if isinstance(o, dict) and o]
        line_qty = int(ln.get("qty") or 0) or 0

        base_name = _safe_sheet_name(model_code or "Innovent Import")
        name = base_name
        n = 2
        while name in sheet_names_seen:
            suffix = f"_{n}"
            name = _safe_sheet_name(base_name[: 31 - len(suffix)] + suffix)
            n += 1
        sheet_names_seen.add(name)

        ws = wb.create_sheet(name)
        _write_one_sheet(
            ws,
            manufacturer=manufacturer,
            equipment="Air Handler",
            model_code=model_code,
            tags=tags,
            header=header,
            line_qty=line_qty,
            options=options,
            output_type=output_type
        )

    return wb


# -----------------------------
# ✅ ROUTE-COMPATIBLE RUNNER
# -----------------------------

def convert_innovent_pdf_to_xlsx_bytes(
    pdf_bytes: bytes,
    job_name: str | None = None,
    output_type: str = "all_in_one",
) -> tuple[bytes, str]:
    """
    Keep same signature as your existing Flask premise:
      converter_fn(pdf_bytes=..., job_name=..., output_type=...)
    NOTE: pdf_bytes is actually "uploaded file bytes" for Innovent (Word bytes).
    """
    job = (job_name or "").strip() or "job"
    parsed = parse_innovent_word_bytes(pdf_bytes)

    wb = write_innovent_workbook(parsed, output_type=output_type)

    out = io.BytesIO()
    wb.save(out)
    xlsx_bytes = out.getvalue()
    out.close()

    filename = f"{job}_innovent_{(output_type or 'all_in_one').strip().lower()}_template_output.xlsx".replace(" ", "_")
    return xlsx_bytes, filename


