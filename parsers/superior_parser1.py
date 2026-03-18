from __future__ import annotations

import io
import re
from typing import Any

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

import logging
logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────
# Shared helpers
# ─────────────────────────────────────────────

MONEY_RE = re.compile(r"\$\s*([0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]{2})?)")


def _money_to_float(s: str) -> float:
    return float(s.replace(",", ""))


def _safe_sheet_name(name: str) -> str:
    bad = r'[:\\/?*\[\]]'
    name = re.sub(bad, "-", (name or "")).strip()
    return (name[:31] or "Sheet1")


# ─────────────────────────────────────────────
# DOCX PARSER
# ─────────────────────────────────────────────

def parse_docx_bytes(file_bytes: bytes, filename: str = "") -> dict[str, Any]:
    """
    Parse a Superior Boiler-style .docx quote into the shared intermediate format.

    Structure of the docx:
      Table 0  – Quote Info (customer name/address)
      Table 1  – Terms & Conditions
      Table 2  – Column headers: ID | Product | Qty | Unit Price | Total Price
      Table 3  – Line items (paired rows: even=part#, odd=description)
      Table 4  – Discount + Total summary

    Row pairs in Table 3:
      Row N (even index 0,2,4,...): ID | PartNumber | Qty | UnitPrice | TotalPrice
      Row N+1 (odd index 1,3,5,...): ID | Description | Qty | UnitPrice | TotalPrice
    """
    from docx import Document  # lazy import

    doc = Document(io.BytesIO(file_bytes))

    # ── Header ──────────────────────────────────
    quote_number = re.sub(r'\.(docx|pdf)$', '', filename, flags=re.IGNORECASE) if filename else ""

    customer = ""
    job_name = ""
    try:
        cell_text = doc.tables[0].rows[1].cells[1].text.strip()
        lines = [l.strip() for l in cell_text.splitlines() if l.strip()]
        if lines:
            customer = lines[0]
        if len(lines) > 1:
            job_name = lines[1]
    except (IndexError, AttributeError):
        pass

    quote_date = ""
    try:
        created = doc.core_properties.created
        if created:
            quote_date = created.strftime("%m/%d/%Y")
    except Exception:
        pass

    header = {
        "quote_number": quote_number,
        "quote_date": quote_date,
        "requested_ship_date": "",
        "job_name": job_name,
        "customer": customer,
    }

    # ── Line items ─────────────────────────────
    items_table = None
    for t in doc.tables:
        if len(t.columns) == 5:
            first_row_cells = [c.text.strip() for c in t.rows[0].cells]
            if first_row_cells[0].isdigit():
                items_table = t
                break

    if items_table is None:
        return {"manufacturer": "Superior Boiler", "header": header, "lines": []}

    data_rows = list(items_table.rows)

    groups: dict[str, list] = {}
    for row in data_rows:
        cells = [c.text.strip() for c in row.cells]
        item_id = cells[0]
        if not item_id:
            continue
        groups.setdefault(item_id, []).append(cells)

    all_items: list[dict[str, Any]] = []
    for item_id in sorted(groups.keys(), key=lambda x: int(x) if x.isdigit() else 0):
        pair = groups[item_id]
        if len(pair) < 2:
            continue
        part_row = pair[0]
        desc_row = pair[1]

        part_number = part_row[1]
        description = desc_row[1]
        try:
            qty = int(part_row[2]) if part_row[2] else 1
        except ValueError:
            qty = 1

        unit_price = 0.0
        total_price = 0.0
        m = MONEY_RE.search(part_row[3])
        if m:
            unit_price = _money_to_float(m.group(1))
        m = MONEY_RE.search(part_row[4])
        if m:
            total_price = _money_to_float(m.group(1))

        all_items.append({
            "item_id": item_id,
            "part_number": part_number,
            "description": description,
            "qty": qty,
            "unit_price": unit_price,
            "total_price": total_price,
        })

    return _build_parsed(all_items, header)


# ─────────────────────────────────────────────
# PDF PARSER (Superior Boiler) - dual format
# ─────────────────────────────────────────────

# Combined format (demo server): one line per item
# e.g. "1 3X2005S15M 1 $ 119,779.36 $119,779.36"
_PDF_ITEM_COMBINED_RE = re.compile(
    r"^(\d{1,3})\s+"
    r"([A-Z0-9][A-Z0-9\-]{2,})\s+"
    r"(\d+)\s+"
    r"\$\s*([\d,]+\.\d{2})\s+"
    r"\$\s*([\d,]+\.\d{2})\s*$"
)

# Split format (new server): each field on its own line
_ID_RE      = re.compile(r"^\s*(\d{1,3})\s*$")
_PART_RE    = re.compile(r"^\s*([A-Z0-9][A-Z0-9\-]{2,})\s*$")
_QTY_RE     = re.compile(r"^\s*(\d+)\s*$")
_PRICE_RE   = re.compile(r"^\s*\$\s*[\d,]+\.\d{2}\s*$")

_PDF_SKIP_RE = re.compile(
    r"^(ID$|Product$|Qty$|Unit Price$|Total Price$|"
    r"Wren Industries|,\s*$|,\s*LA|P a g e|Quote Information|"
    r"Terms\s*&\s*Conditions|QUOTES:|PAYMENT TERMS:|FREIGHT CHARGES:|"
    r"This Proposal|By:|Title:|Reference P\.O\.|As Quoted|With Exception:|"
    r"Discounts|Applied:|Total:|Superior Boiler|Prepared for|Reviewed for|"
    r"Our current lead time|Submittal Approval|Immediate Release)",
    re.IGNORECASE
)


def _detect_pdf_format(lines: list[str]) -> str:
    for line in lines:
        if _PDF_ITEM_COMBINED_RE.match(line.strip()):
            return "combined"
    return "split"


def _parse_pdf_combined(lines: list[str]) -> list[dict[str, Any]]:
    items = []
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].strip()
        i += 1
        if not line or _PDF_SKIP_RE.match(line):
            continue
        m = _PDF_ITEM_COMBINED_RE.match(line)
        if not m:
            continue
        item_id    = m.group(1)
        part_num   = m.group(2)
        qty        = int(m.group(3))
        unit_price = _money_to_float(m.group(4))
        total      = _money_to_float(m.group(5))
        description = ""
        while i < n:
            nxt = lines[i].strip()
            i += 1
            if not nxt: continue
            if _PDF_SKIP_RE.match(nxt): continue
            if _PDF_ITEM_COMBINED_RE.match(nxt):
                i -= 1
                break
            description = nxt
            while i < n:
                cont = lines[i].strip()
                if not cont or _PDF_ITEM_COMBINED_RE.match(cont) or _PDF_SKIP_RE.match(cont):
                    break
                description += " " + cont
                i += 1
            break
        items.append({"item_id": item_id, "part_number": part_num,
                      "description": description, "qty": qty,
                      "unit_price": unit_price, "total_price": total})
    return items


def _find_pdf_blocks(lines: list[str]) -> list[tuple]:
    """Find (id_idx, part_idx, qty_idx, p1_idx, p2_idx) for each item in split format."""
    blocks = []
    n = len(lines)
    for i in range(n - 4):
        if (_ID_RE.match(lines[i]) and
                _PART_RE.match(lines[i + 1]) and
                _QTY_RE.match(lines[i + 2]) and
                _PRICE_RE.match(lines[i + 3]) and
                _PRICE_RE.match(lines[i + 4])):
            blocks.append((i, i + 1, i + 2, i + 3, i + 4))
    return blocks


def _parse_pdf_split(lines: list[str]) -> list[dict[str, Any]]:
    blocks = _find_pdf_blocks(lines)
    if not blocks:
        return []

    items = []
    n = len(lines)

    for b_idx, (id_i, pn_i, qty_i, p1_i, p2_i) in enumerate(blocks):
        item_id    = lines[id_i].strip()
        part_num   = lines[pn_i].strip()
        qty        = int(lines[qty_i].strip())
        unit_price = _money_to_float(MONEY_RE.search(lines[p1_i]).group(1))
        total      = _money_to_float(MONEY_RE.search(lines[p2_i]).group(1))

        # Description = lines between p2_i and next block start
        next_start = blocks[b_idx + 1][0] if b_idx + 1 < len(blocks) else n
        desc_lines = []
        for k in range(p2_i + 1, next_start):
            l = lines[k].strip()
            if not l or _PDF_SKIP_RE.match(l):
                continue
            # Stop if we hit discount/total lines
            if MONEY_RE.match(l) and not l.startswith('$'):
                break
            # Skip bare dollar amounts (discount totals bleeding in)
            if re.match(r"^\$[\d,]+\.\d{2}$", l):
                continue
            desc_lines.append(l)

        desc = " ".join(desc_lines).strip()
        items.append({"item_id": item_id, "part_number": part_num,
                      "description": desc, "qty": qty,
                      "unit_price": unit_price, "total_price": total})
    return items


def parse_pdf_bytes(pdf_bytes: bytes, filename: str = "") -> dict[str, Any]:
    """
    Parse a Superior Boiler-style .pdf quote.
    Auto-detects combined (old server) vs split (new server) text extraction format.
    """
    try:
        import fitz
    except ImportError:
        raise ImportError("PyMuPDF (fitz) is required. Install with: pip install PyMuPDF")

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    # ── Header from page 1 ──────────────────────
    p1 = doc[0].get_text("text") if len(doc) > 0 else ""

    quote_number = ""
    m = re.search(r"Quote\s*Number:\s*([0-9A-Za-z\-]+)", p1)
    if m: quote_number = m.group(1).strip()

    quote_date = ""
    m = re.search(r"Date\s*of\s*Quote:\s*([\d/]+)", p1)
    if m: quote_date = m.group(1).strip()

    expiry_date = ""
    m = re.search(r"Expiration\s*Date:\s*([\d/]+)", p1)
    if m: expiry_date = m.group(1).strip()

    customer = ""
    job_name = ""
    m = re.search(r"Quote\s*Information\s*\n(.+?)\n(.+?)\n", p1)
    if m:
        customer = m.group(1).strip()
        job_name = m.group(2).strip()

    header = {
        "quote_number": quote_number,
        "quote_date": quote_date,
        "expiry_date": expiry_date,
        "requested_ship_date": "",
        "job_name": job_name,
        "customer": customer,
    }

    # ── Line items from pages 2+ ────────────────
    all_lines: list[str] = []
    for page_num in range(1, len(doc)):
        all_lines.extend((doc[page_num].get_text("text") or "").splitlines())

    fmt = _detect_pdf_format(all_lines)
    logger.info(f"superior_parser PDF: detected format='{fmt}' from {filename}")

    if fmt == "combined":
        all_items = _parse_pdf_combined(all_lines)
    else:
        all_items = _parse_pdf_split(all_lines)

    logger.info(f"superior_parser PDF: parsed {len(all_items)} items")
    return _build_parsed(all_items, header)


# ─────────────────────────────────────────────
# Shared builder
# ─────────────────────────────────────────────

def _build_parsed(all_items: list[dict[str, Any]], header: dict[str, Any]) -> dict[str, Any]:
    if not all_items:
        return {"manufacturer": "Superior Boiler", "header": header, "lines": []}

    main = all_items[0]
    include_items = all_items[1:]

    options = [
        {
            "part_number": it["part_number"],
            "desc": it["description"],
            "qty": it["qty"],
            "unit_price": it["unit_price"],
            "total_price": it["total_price"],
            "tag": it["item_id"],
        }
        for it in include_items
    ]

    full_desc = main["description"]
    equipment = full_desc.split(",")[0].strip() if "," in full_desc else full_desc

    line = {
        "model_code": main["part_number"],
        "description": full_desc,
        "equipment": equipment,
        "tagging": "",
        "tags": [],
        "qty": main["qty"],
        "list_each": main["unit_price"],
        "total_list": main["total_price"],
        "options": options,
    }

    return {
        "manufacturer": "Superior Boiler",
        "header": header,
        "lines": [line],
    }


# ─────────────────────────────────────────────
# UNIFIED ENTRY POINT
# ─────────────────────────────────────────────

def parse_file_bytes(file_bytes: bytes, filename: str = "") -> dict[str, Any]:
    """
    Detect file type by extension or magic bytes and dispatch to the right parser.
    """
    fname_lower = (filename or "").lower()

    if fname_lower.endswith(".docx"):
        try:
            return parse_docx_bytes(file_bytes, filename=filename)
        except Exception as e:
            logger.warning("DOCX parse failed (%s), trying PDF fallback", e)
            return parse_pdf_bytes(file_bytes, filename=filename)

    if fname_lower.endswith(".pdf"):
        return parse_pdf_bytes(file_bytes, filename=filename)

    # No extension — try docx magic bytes (PK zip signature)
    if file_bytes[:2] == b'PK':
        try:
            return parse_docx_bytes(file_bytes, filename=filename)
        except Exception as e:
            logger.warning("DOCX parse failed (%s), trying PDF fallback", e)
            return parse_pdf_bytes(file_bytes, filename=filename)

    return parse_pdf_bytes(file_bytes, filename=filename)


# ─────────────────────────────────────────────
# XLSX WRITER
# ─────────────────────────────────────────────

def _write_one_sheet(
    ws,
    *,
    manufacturer: str,
    equipment: str,
    model_code: str,
    description: str,
    tags: list[str],
    header: dict[str, Any],
    options: list[dict[str, Any]],
) -> None:
    top_headers = ["Equipment", "Manufacturer", "Model", "Part Number",
                   "Description (Not Overwritten)", "Notes (Not Overwritten)"]
    for col, h in enumerate(top_headers, start=1):
        ws.cell(row=1, column=col, value=h)

    model_short = model_code.split("-")[0] if model_code else ""
    quote_no    = header.get("quote_number", "")
    job_name    = header.get("job_name", "")
    quote_date  = header.get("quote_date", "")
    ship_date   = header.get("requested_ship_date", "")
    customer    = header.get("customer", "")
    expiry_date = header.get("expiry_date", "")

    ws.cell(row=2, column=1, value=equipment)
    ws.cell(row=2, column=2, value=manufacturer)
    ws.cell(row=2, column=3, value=model_short)
    ws.cell(row=2, column=4, value=model_code)
    ws.cell(row=2, column=5, value=description)

    notes_parts: list[str] = []
    if quote_no:    notes_parts.append(f"Quote #: {quote_no}")
    if job_name:    notes_parts.append(f"Job: {job_name}")
    if customer:    notes_parts.append(f"Customer: {customer}")
    if quote_date:  notes_parts.append(f"Quote Date: {quote_date}")
    if expiry_date: notes_parts.append(f"Expiry: {expiry_date}")
    if ship_date:   notes_parts.append(f"Ship Date: {ship_date}")
    ws.cell(row=2, column=6, value="\n".join(notes_parts))

    opt_headers = [
        "Tag", "Part Number", "Feature", "Description",
        "Qty", "List Price", "LP Ext.", "Buy Mult.", "Net Price",
        "Markup", "Margin", "Sell Price", "Weight", "Freight",
        "Fr. Multi.", "Alignment", "Subtotal", "Option Price",
    ]
    start_col = 3
    for i, h in enumerate(opt_headers):
        ws.cell(row=3, column=start_col + i, value=h)

    row_idx = 4
    for opt in (options or []):
        part_num   = opt.get("part_number", "")
        desc       = opt.get("desc", "")
        unit_price = float(opt.get("unit_price") or 0)
        opt_qty    = int(opt.get("qty") or 1)
        lp_ext     = unit_price * opt_qty
        tag        = opt.get("tag") or (", ".join(tags) if tags else "")

        ws.cell(row=row_idx, column=3,  value=tag)
        ws.cell(row=row_idx, column=4,  value=part_num)
        ws.cell(row=row_idx, column=5,  value="Include")
        ws.cell(row=row_idx, column=6,  value=desc)
        ws.cell(row=row_idx, column=7,  value=opt_qty)
        ws.cell(row=row_idx, column=8,  value=unit_price)
        ws.cell(row=row_idx, column=9,  value=lp_ext)
        for c in range(10, 21):
            ws.cell(row=row_idx, column=c, value=0)
        ws.cell(row=row_idx, column=20, value=unit_price)
        row_idx += 1

    header_fill = PatternFill("solid", fgColor="D9EAD3")
    top_font    = Font(bold=True, italic=True, name="Arial")
    opt_font    = Font(bold=True, name="Arial")
    body_font   = Font(name="Arial")

    for c in range(1, 7):
        cell = ws.cell(row=1, column=c)
        cell.font = top_font; cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for c in range(3, 21):
        cell = ws.cell(row=3, column=c)
        cell.font = opt_font; cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for r in range(2, max(row_idx, 4)):
        for c in range(1, 21):
            cell = ws.cell(row=r, column=c)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.font = body_font

    widths = {1:18, 2:18, 3:16, 4:20, 5:45, 6:45,
              7:8, 8:12, 9:12, 10:10, 11:12, 12:10,
              13:10, 14:12, 15:10, 16:10, 17:10, 18:10, 19:12, 20:12}
    for col, w in widths.items():
        ws.column_dimensions[get_column_letter(col)].width = w

    ws.freeze_panes = "A4"


def write_template_workbook(parsed: dict[str, Any]) -> Workbook:
    wb = Workbook()
    wb.remove(wb.active)

    manufacturer = parsed.get("manufacturer", "")
    header       = parsed.get("header", {})
    lines        = parsed.get("lines", [])

    if not lines:
        ws = wb.create_sheet("Import")
        _write_one_sheet(ws, manufacturer=manufacturer, equipment="",
                         model_code="", description="", tags=[],
                         header=header, options=[])
        return wb

    sheet_names_seen: set[str] = set()

    for ln in lines:
        model_code  = ln.get("model_code", "") or ""
        description = ln.get("description", "") or ""
        tagging     = ln.get("tagging", "") or ""
        tags        = ln.get("tags") or ([tagging] if tagging else [])
        options     = ln.get("options", []) or []
        equipment   = ln.get("equipment", "")

        base_name = _safe_sheet_name(model_code or "Import")
        name = base_name
        n = 2
        while name in sheet_names_seen:
            suffix = f"_{n}"
            name = _safe_sheet_name(base_name[:31 - len(suffix)] + suffix)
            n += 1
        sheet_names_seen.add(name)

        ws = wb.create_sheet(title=name)
        _write_one_sheet(
            ws,
            manufacturer=manufacturer,
            equipment=equipment,
            model_code=model_code,
            description=description,
            tags=tags,
            header=header,
            options=options,
        )

    return wb


# ─────────────────────────────────────────────
# PUBLIC API
# ─────────────────────────────────────────────

def convert_superior_to_xlsx_bytes(
    pdf_bytes: bytes,
    filename: str = "",
    job_name: str | None = None,
    output_type: str | None = None,
    **kwargs,
) -> tuple[bytes, str]:
    """
    Main entry point. Accepts .docx or .pdf bytes.
    Returns (xlsx_bytes, suggested_filename).
    """
    parsed = parse_file_bytes(pdf_bytes, filename=filename)

    if job_name:
        parsed["header"]["job_name"] = job_name.strip()

    wb = write_template_workbook(parsed)

    out = io.BytesIO()
    wb.save(out)
    xlsx_bytes = out.getvalue()
    out.close()

    job = (parsed["header"].get("job_name") or
           re.sub(r'\.(docx|pdf)$', '', filename, flags=re.IGNORECASE) or
           "output")
    return xlsx_bytes, f"{job}_template_output.xlsx"


# ─────────────────────────────────────────────
# Quick test when run directly
# ─────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    path = sys.argv[1] if len(sys.argv) > 1 else "/mnt/user-data/uploads/04040066-2.pdf"
    with open(path, "rb") as f:
        data = f.read()
    fname = path.split("/")[-1]
    xlsx_bytes, out_name = convert_superior_to_xlsx_bytes(data, filename=fname)
    out_path = f"/mnt/user-data/outputs/{out_name}"
    with open(out_path, "wb") as f:
        f.write(xlsx_bytes)
    print(f"Written: {out_path}")